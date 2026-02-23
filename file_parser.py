# ============================================================
# file_parser.py  —  Quality Intelligence Engine v3
# Universal file parser — CSV, PDF, TXT, Word (.docx)
# Returns a pandas DataFrame regardless of input format.
# ============================================================

import pandas as pd
import numpy as np
import io
import re
from typing import Tuple


def parse_file(uploaded_file) -> Tuple[pd.DataFrame | None, str, str]:
    """
    Main entry point. Takes a Streamlit uploaded file object.
    Returns: (dataframe, file_type, message)
    file_type: 'csv', 'pdf', 'txt', 'docx', 'excel'
    message: success info or error description
    """
    name = uploaded_file.name.lower()
    raw  = uploaded_file.read()

    if name.endswith(".csv"):
        return _parse_csv(raw)
    elif name.endswith(".pdf"):
        return _parse_pdf(raw)
    elif name.endswith(".txt"):
        return _parse_txt(raw)
    elif name.endswith(".docx") or name.endswith(".doc"):
        return _parse_word(raw)
    elif name.endswith(".xlsx") or name.endswith(".xls"):
        return _parse_excel(raw)
    else:
        return None, "unknown", f"Unsupported file type: {name.split('.')[-1]}"


# ── CSV ───────────────────────────────────────────────────────────────────────
def _parse_csv(raw: bytes) -> Tuple:
    try:
        # Try common encodings
        for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
            try:
                df = pd.read_csv(io.BytesIO(raw), encoding=enc)
                df = _clean_df(df)
                return df, "csv", f"CSV loaded: {len(df):,} rows × {len(df.columns)} columns"
            except UnicodeDecodeError:
                continue
        return None, "csv", "Could not decode CSV — try saving as UTF-8"
    except Exception as e:
        return None, "csv", f"CSV parse error: {e}"


# ── Excel ─────────────────────────────────────────────────────────────────────
def _parse_excel(raw: bytes) -> Tuple:
    try:
        xf = pd.ExcelFile(io.BytesIO(raw))
        sheets = xf.sheet_names

        if len(sheets) == 1:
            df = pd.read_excel(io.BytesIO(raw), sheet_name=0)
            df = _clean_df(df)
            return df, "excel", f"Excel loaded (sheet: {sheets[0]}): {len(df):,} rows × {len(df.columns)} columns"
        else:
            # Pick the sheet with the most data
            best_df   = None
            best_rows = 0
            best_name = ""
            for sheet in sheets:
                try:
                    tmp = pd.read_excel(io.BytesIO(raw), sheet_name=sheet)
                    if len(tmp) > best_rows:
                        best_df   = tmp
                        best_rows = len(tmp)
                        best_name = sheet
                except Exception:
                    continue
            if best_df is not None:
                best_df = _clean_df(best_df)
                return best_df, "excel", f"Excel loaded (largest sheet: '{best_name}'): {len(best_df):,} rows × {len(best_df.columns)} columns"
            return None, "excel", "Could not read any sheet"
    except Exception as e:
        return None, "excel", f"Excel parse error: {e}"


# ── PDF ───────────────────────────────────────────────────────────────────────
def _parse_pdf(raw: bytes) -> Tuple:
    try:
        import pdfplumber
    except ImportError:
        return None, "pdf", "pdfplumber not installed. Run: pip install pdfplumber"

    try:
        all_tables = []
        all_text   = []

        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                # Try to extract tables first
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        all_tables.append(table)
                # Also grab text for fallback
                text = page.extract_text()
                if text:
                    all_text.append(text)

        # If tables found — convert to DataFrame
        if all_tables:
            df = _tables_to_df(all_tables)
            if df is not None and len(df) > 0:
                df = _clean_df(df)
                return df, "pdf", f"PDF table extracted: {len(df):,} rows × {len(df.columns)} columns"

        # Fallback: parse text as structured data
        if all_text:
            full_text = "\n".join(all_text)
            df = _text_to_df(full_text)
            if df is not None and len(df) > 0:
                df = _clean_df(df)
                return df, "pdf", f"PDF text parsed: {len(df):,} rows × {len(df.columns)} columns"

        return None, "pdf", "No tables or structured data found in PDF. Try exporting as CSV from your QA tool."

    except Exception as e:
        return None, "pdf", f"PDF parse error: {e}"


# ── Word (.docx) ──────────────────────────────────────────────────────────────
def _parse_word(raw: bytes) -> Tuple:
    try:
        from docx import Document
    except ImportError:
        return None, "docx", "python-docx not installed. Run: pip install python-docx"

    try:
        doc     = Document(io.BytesIO(raw))
        tables  = doc.tables
        all_text = "\n".join([p.text for p in doc.paragraphs])

        if tables:
            # Convert Word tables to DataFrame
            all_dfs = []
            for table in tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                if rows:
                    # First row as header
                    headers = rows[0]
                    data    = rows[1:]
                    if data:
                        try:
                            df = pd.DataFrame(data, columns=headers)
                            all_dfs.append(df)
                        except Exception:
                            pass

            if all_dfs:
                # Pick largest table
                best = max(all_dfs, key=len)
                best = _clean_df(best)
                return best, "docx", f"Word table extracted: {len(best):,} rows × {len(best.columns)} columns"

        # Fallback: parse text content
        if all_text.strip():
            df = _text_to_df(all_text)
            if df is not None and len(df) > 0:
                df = _clean_df(df)
                return df, "docx", f"Word text parsed: {len(df):,} rows × {len(df.columns)} columns"

        return None, "docx", "No tables found in Word document. Try copy-pasting data into a CSV."

    except Exception as e:
        return None, "docx", f"Word parse error: {e}"


# ── TXT ───────────────────────────────────────────────────────────────────────
def _parse_txt(raw: bytes) -> Tuple:
    try:
        for enc in ["utf-8", "utf-8-sig", "latin-1"]:
            try:
                text = raw.decode(enc)
                break
            except UnicodeDecodeError:
                continue
        else:
            return None, "txt", "Could not decode text file"

        # Try CSV-like parsing first (tab or comma separated)
        for sep in [",", "\t", "|", ";"]:
            try:
                df = pd.read_csv(io.StringIO(text), sep=sep)
                if len(df.columns) >= 2 and len(df) >= 1:
                    df = _clean_df(df)
                    return df, "txt", f"TXT parsed as delimited: {len(df):,} rows × {len(df.columns)} columns"
            except Exception:
                continue

        # Try structured text parsing
        df = _text_to_df(text)
        if df is not None and len(df) > 0:
            df = _clean_df(df)
            return df, "txt", f"TXT text parsed: {len(df):,} rows × {len(df.columns)} columns"

        return None, "txt", "TXT file could not be parsed as tabular data. Use CSV format for best results."

    except Exception as e:
        return None, "txt", f"TXT parse error: {e}"


# ── Helpers ───────────────────────────────────────────────────────────────────
def _tables_to_df(tables: list) -> pd.DataFrame | None:
    """Convert list of raw table arrays to a single DataFrame."""
    all_dfs = []
    for table in tables:
        if not table or len(table) < 2:
            continue
        headers = [str(h).strip() if h else f"col_{i}" for i, h in enumerate(table[0])]
        rows    = []
        for row in table[1:]:
            rows.append([str(c).strip() if c else "" for c in row])
        try:
            df = pd.DataFrame(rows, columns=headers)
            all_dfs.append(df)
        except Exception:
            continue
    if not all_dfs:
        return None
    return max(all_dfs, key=len)


def _text_to_df(text: str) -> pd.DataFrame | None:
    """
    Try to parse unstructured text into a DataFrame.
    Handles: key-value pairs, space-aligned columns, colon-separated data.
    """
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if not lines:
        return None

    # Try: detect header row by looking for consistent column structure
    # Find lines with similar number of tokens (likely table rows)
    token_counts = [len(re.split(r"\s{2,}|\t|,|\|", l)) for l in lines]
    most_common  = pd.Series(token_counts).mode()
    if not most_common.empty and most_common[0] >= 2:
        target_cols = most_common[0]
        table_lines = [l for l, tc in zip(lines, token_counts) if tc == target_cols]
        if len(table_lines) >= 2:
            sep = r"\s{2,}|\t|,|\|"
            rows = [re.split(sep, l) for l in table_lines]
            try:
                df = pd.DataFrame(rows[1:], columns=rows[0])
                if len(df) > 0 and len(df.columns) >= 2:
                    return df
            except Exception:
                pass

    # Last resort: try key:value pairs → one row per entry
    kv_rows = {}
    for line in lines:
        if ":" in line:
            parts = line.split(":", 1)
            key   = parts[0].strip().lower().replace(" ", "_")
            val   = parts[1].strip()
            if key and val:
                kv_rows[key] = val
    if len(kv_rows) >= 3:
        return pd.DataFrame([kv_rows])

    return None


def _clean_df(df: pd.DataFrame) -> pd.DataFrame:
    """Standardise column names and fix data types."""
    # Normalise column names
    df.columns = (df.columns.astype(str)
                  .str.strip()
                  .str.lower()
                  .str.replace(r"[\s\-/\\]+", "_", regex=True)
                  .str.replace(r"[^\w]", "", regex=True)
                  .str.strip("_"))

    # Remove fully empty rows and columns
    df = df.dropna(how="all").reset_index(drop=True)
    df = df.loc[:, df.notna().any()]

    # Auto-convert numeric columns
    for col in df.columns:
        try:
            converted = pd.to_numeric(df[col], errors="coerce")
            if converted.notna().sum() / max(len(df), 1) > 0.5:
                df[col] = converted
        except Exception:
            pass

    return df


def get_file_type_label(file_type: str) -> str:
    labels = {
        "csv":   "📊 CSV Spreadsheet",
        "pdf":   "📄 PDF Document",
        "txt":   "📝 Text File",
        "docx":  "📋 Word Document",
        "excel": "📊 Excel Spreadsheet",
    }
    return labels.get(file_type, "📁 File")
